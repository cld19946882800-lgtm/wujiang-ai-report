import streamlit as st
import pandas as pd
import io
import os
import openai
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib import rcParams

# 修复PDF中文乱码、空白核心配置
rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei', 'PingFang SC']
rcParams['axes.unicode_minus'] = False
rcParams['pdf.fonttype'] = 42
rcParams['ps.fonttype'] = 42
rcParams['figure.dpi'] = 150

VOLC_API_KEY = "ark-2358e229-f997-477c-a223-f71a3c1d67f4-13835"
VOLC_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
VOLC_MODEL_NAME = "ep-20260513185957-wfw8z"

DEEPSEEK_API_KEY = "sk-0d889790af4d4eddbb912030535b49a2"
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"
DEEPSEEK_MODEL_NAME = "deepseek-chat"

DEFAULT_DOMAINS = {
    "知识产权": {
        "name": "知识产权",
        "icon": "📈",
        "description": "专利申请、授权、失效、PCT等数据分析",
        "required_files": [
            {"key": "f_valid", "label": "《有效专利清单》", "desc": "包含专利号、专利权人、申请日、授权日、专利类型、法律状态、IPC分类号、申请人类型、所属区域等"},
            {"key": "f_pct", "label": "《PCT申请清单》", "desc": "PCT国际申请记录"},
            {"key": "f_auth", "label": "《发明授权清单》", "desc": "本月新增授权专利"},
            {"key": "f_loss", "label": "《失效专利清单》", "desc": "本月失效专利清单"}
        ],
        "core_fields": {
            "group_by_field": "所属区域",
            "patent_type_field": "专利类型",
            "legal_status_field": "法律状态",
            "ipc_field": "IPC分类号",
            "applicant_type_field": "申请人类型",
            "apply_date_field": "申请日",
            "auth_date_field": "授权日",
            "loss_date_field": "失效日",
            "applicant_name_field": "专利权人名称"
        },
        "metrics": [
            {"key": "有效专利总量", "source": "f_valid", "label": "有效专利总量(件)"},
            {"key": "有效发明专利存量", "source": "f_valid", "label": "有效发明专利存量(件)"},
            {"key": "本月新增发明专利授权", "source": "f_auth", "label": "本月新增发明专利授权(件)"},
            {"key": "本月PCT国际申请", "source": "f_pct", "label": "本月PCT国际申请(件)"},
            {"key": "本月失效专利", "source": "f_loss", "label": "本月失效专利(件)"},
            {"key": "发明专利占比", "source": "f_valid", "label": "发明专利占比(%)"},
            {"key": "万人发明专利拥有量", "source": "f_valid", "label": "万人发明专利拥有量(件/万人)"}
        ],
        "title": "吴江区知识产权统计深度简报",
        "prompt_template": """你是吴江区市场监管局（知识产权局）分析师，严格按照政府正式「本期研判」模板生成，语言精炼、结论化、无套话，字数控制在200字左右。

内容必须包含：
1. 吴江区全区域专利总量、有效发明专利存量；
2. 本月新增授权、PCT申请、失效专利数量；
3. 吴江区所有下属区域的专利数量明细；
4. 授权领先主体TOP1；
5. 重点技术领域TOP5。

格式：
标题：吴江区知识产权_本期研判
正文：吴江区全区域专利总量{吴江区全区域专利总量}件，有效发明专利存量{有效发明专利存量}件。本月新增发明专利授权{本月新增发明专利授权}件，PCT国际申请{本月PCT国际申请}件，失效专利{本月失效专利}件。各区域专利数量：{吴江区所有区域专利明细}。{top_applicant}授权领先。重点技术领域集中在{top_ipc}。

不要分段、不要小标题、不要废话、不要修饰词、不要建议、不要风险。

【数据】
吴江区全区域专利总量：{吴江区全区域专利总量}
有效发明专利存量：{有效发明专利存量}
本月新增发明专利授权：{本月新增发明专利授权}
本月PCT国际申请：{本月PCT国际申请}
本月失效专利：{本月失效专利}
吴江区所有区域专利明细：{吴江区所有区域专利明细}
top_applicant：{top_applicant}
top_ipc：{top_ipc}
"""
    },
    "投诉举报": {
        "name": "投诉举报",
        "icon": "📞",
        "description": "12345投诉举报数据分析",
        "required_files": [
            {"key": "f_complaint", "label": "《投诉清单》", "desc": "消费者投诉记录"},
            {"key": "f_report", "label": "《举报清单》", "desc": "违法举报记录"},
            {"key": "f_resolve", "label": "《办结清单》", "desc": "已办结案件清单"},
            {"key": "f_satisfaction", "label": "《满意度评价》", "desc": "群众满意度评价"}
        ],
        "group_by_field": "投诉地区",
        "metrics": [
            {"key": "投诉量", "source": "f_complaint", "label": "投诉量(件)"},
            {"key": "举报量", "source": "f_report", "label": "举报量(件)"},
            {"key": "办结量", "source": "f_resolve", "label": "办结量(件)"},
            {"key": "满意度", "source": "f_satisfaction", "label": "满意度(%)"}
        ],
        "top_field": "投诉地区",
        "type_field": "投诉类型",
        "class_field": "行业分类",
        "reason_field": "处理状态",
        "title": "吴江区投诉举报情况分析简报",
        "prompt_template": """你是一位吴江区市场监管局的高级分析师。请根据以下数据，写一段高质量的《投诉举报月度分析简报》。
本月共受理投诉 {投诉量} 件，举报 {举报量} 件。
办结案件 {办结量} 件，办结率较高。
热点投诉领域主要集中在【{top_types}】。
区域分布上，【{top3_districts}】投诉举报量较多。
要求：用"政府智库"的口吻，分两个自然段，总字数300-400字左右。"""
    },
    "食品安全": {
        "name": "食品安全",
        "icon": "🍱",
        "description": "食品安全监管、抽检、处罚数据分析",
        "required_files": [
            {"key": "f_inspection", "label": "《监督检查清单》", "desc": "日常检查记录"},
            {"key": "f_sampling", "label": "《抽检结果清单》", "desc":  "食品安全抽检结果"},
            {"key": "f_penalty", "label": "《行政处罚清单》", "desc": "行政处罚记录"},
            {"key": "f_recall", "label": "《产品召回清单》", "desc": "问题产品召回记录"}
        ],
        "group_by_field": "区域",
        "metrics": [
            {"key": "检查次数", "source": "f_inspection", "label": "检查次数(次)"},
            {"key": "抽检批次", "source": "f_sampling", "label": "抽检批次(批)"},
            {"key": "处罚案件", "source": "f_penalty", "label": "处罚案件(件)"},
            {"key": "召回数量", "source": "f_recall", "label": "召回数量(个)"}
        ],
        "top_field": "经营主体",
        "type_field": "业态类型",
        "class_field": "食品类别",
        "reason_field": "问题类型",
        "title": "吴江区食品安全监管情况分析简报",
        "prompt_template": """你是一位吴江区市场监管局的高级食品安全分析师。请根据以下数据，写一段高质量的《食品安全月度监管简报》。
本月共开展监督检查 {检查次数} 次，完成抽检 {抽检批次} 批次。
查处行政处罚案件 {处罚案件} 件，问题产品召回 {召回数量} 个。
监管重点领域主要集中在【{top_types}】。
区域分布上，【{top3_districts}】监管任务较重。
要求：用"政府智库"的口吻，分两个自然段，总字数300-400字左右。"""
    },
    "医疗器械": {
        "name": "医疗器械",
        "icon": "🏥",
        "description": "医疗器械生产、经营、监管数据分析",
        "required_files": [
            {"key": "f_manufacture", "label": "《生产企业清单》", "desc": "生产企业台账"},
            {"key": "f_operation", "label": "《经营企业清单》", "desc": "经营企业台账"},
            {"key": "f_inspection", "label": "《监督检查记录》", "desc": "日常检查记录"},
            {"key": "f_adverse", "label": "《不良事件报告》", "desc": "医疗器械不良事件"}
        ],
        "group_by_field": "区域",
        "metrics": [
            {"key": "生产企业", "source": "f_manufacture", "label": "生产企业(家)"},
            {"key": "经营企业", "source": "f_operation", "label": "经营企业(家)"},
            {"key": "检查次数", "source": "f_inspection", "label": "检查次数(次)"},
            {"key": "不良事件", "source": "f_adverse", "label": "不良事件(例)"}
        ],
        "top_field": "企业名称",
        "type_field": "企业类型",
        "class_field": "产品分类",
        "reason_field": "问题类别",
        "title": "吴江区医疗器械监管情况分析简报",
        "prompt_template": """你是一位吴江区市场监管局的高级医疗器械监管分析师。请根据以下数据，写一段高质量的《医疗器械月度监管简报》。
全区现有医疗器械生产企业 {生产} 家，经营企业 {经营_value} 家。
本月开展监督检查 {检查次数} 次，收到不良事件报告 {不良事件} 例。
监管重点企业主要集中在【{top_types}】。
区域分布上，【{top3_districts}】企业数量较多。
要求：用"政府智库"的口吻，分两个自然段，总字数300-400字左右。"""
    }
}

def init_custom_domains():
    if 'custom_domains' not in st.session_state:
        st.session_state.custom_domains = {}

def save_custom_domain(name, files_info):
    st.session_state.custom_domains[name] = {
        "name": name,
        "icon": "📋",
        "description": f"自定义分析领域：{name}",
        "required_files": files_info,
        "title": f"{name}分析报告"
    }

init_custom_domains()
DEFAULT_DOMAINS["自定义"] = {
    "name": "自定义",
    "icon": "📋",
    "description": "自定义上传数据并分析",
    "required_files": [],
    "title": "自定义数据分析报告",
    "prompt_template": """你是一位数据分析专家。请根据用户上传的业务数据进行分析，并撰写一份专业的分析简报。

【数据概况】
数据文件数：{数据文件数}
总记录数：{总记录数}

请根据以上数据：
1. 识别数据中的关键指标和趋势
2. 分析主要特征和模式
3. 用"政府智库"的专业口吻撰写一份300-400字的月度分析简报
4. 简报应包含数据解读、问题发现、建议意见三个部分"""
}
DOMAINS = {**DEFAULT_DOMAINS, **st.session_state.custom_domains}

if 'started' not in st.session_state:
    st.session_state.started = False
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = "豆包大模型 (Volcengine)"
if 'selected_domain' not in st.session_state:
    st.session_state.selected_domain = "知识产权"

st.set_page_config(page_title="吴江区市场监管数据分析平台", layout="wide", page_icon="📊")

st.markdown("""
    <style>
    /* 科技网格背景 - 使用 fixed 定位确保不被遮挡 */
    html, body, .stApp, [data-testid="stAppViewContainer"] {
        background-image: 
            linear-gradient(to right, #e2e8f0 1px, transparent 1px),
            linear-gradient(to bottom, #e2e8f0 1px, transparent 1px) !important;
        background-size: 20px 30px !important;
        background-attachment: fixed !important;
        background-color: #f8fafc !important;
    }
    
    /* 内容区域半透明 */
    .main .block-container {
        background: rgba(255,255,255,0.95);
        border-radius: 10px;
        padding: 20px;
    }

    .flat-title {
        border: 2px solid #e2e8f0; background-color: #ffffff; color: #1e293b;
        font-weight: 900; font-size: 20px; text-align: center;
        padding: 15px; border-radius: 8px; margin: 0 0 20px 0;
        width: 100%; box-sizing: border-box; letter-spacing: 2px;
    }

    .stButton > button {
        box-shadow: 0px 5px 0px 0px #94a3b8, 0px 10px 15px rgba(0,0,0,0.1) !important;
        transform: translateY(0px) !important; transition: all 0.1s !important;
        border-radius: 10px !important; font-weight: bold !important;
    }
    .stButton > button:active {
        box-shadow: 0px 0px 0px 0px #94a3b8 !important; transform: translateY(5px) !important;
    }
    .stButton > button[kind="primary"] {
        box-shadow: 0px 5px 0px 0px #b91c1c, 0px 10px 15px rgba(0,0,0,0.1) !important;
        background-color: #ef4444 !important; color: white !important;
    }
    .stButton > button[kind="primary"]:active {
        box-shadow: 0px 0px 0px 0px #b91c1c !important; transform: translateY(5px) !important;
    }

    [data-testid="stFileUploader"] {
        background-color: white; border: 1px solid #e2e8f0; border-radius: 8px; padding: 10px; margin-bottom: 15px;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📊 吴江区市场监管数据分析平台")
st.markdown("选择分析领域，上传业务数据，AI 自动生成多形态分析简报。")
st.markdown("---")

domain_list = [f"{DOMAINS[k]['icon']} {DOMAINS[k]['name']}" for k in DOMAINS.keys()]
domain_keys = list(DOMAINS.keys())

def get_uploaded_files(domain_key):
    files = {}
    domain_config = DOMAINS[domain_key]
    for file_config in domain_config["required_files"]:
        files[file_config["key"]] = None
    return files

def process_domain_data(domain_key, uploaded_files):
    dfs = {}
    for key, file_obj in uploaded_files.items():
        if file_obj is not None:
            try:
                df = pd.read_excel(file_obj, dtype=str, keep_default_na=False)
                df = df.fillna('').replace(r'^\s*$', '', regex=True)
                if not df.empty:
                    dfs[key] = df
            except Exception as e:
                print(f"读取文件 {file_obj.name} 出错: {str(e)}")
                continue
    
    all_dfs = [df for df in dfs.values() if not df.empty]
    combined_df = pd.DataFrame()
    if all_dfs:
        all_cols = sorted(list(set().union(*[df.columns for df in all_dfs])))
        normalized_dfs = []
        for df in all_dfs:
            for col in all_cols:
                if col not in df.columns:
                    df[col] = ''
            normalized_dfs.append(df[all_cols])
        combined_df = pd.concat(normalized_dfs, ignore_index=True)
        combined_df = combined_df.fillna('').replace(r'^\s*$', '', regex=True)
    
    if combined_df.empty:
        return {}, pd.DataFrame(), {}, {}, {}, {}, {}, {}, {}

    all_cols = combined_df.columns.tolist()
    totals = {"文件数": len(dfs), "总记录数": len(combined_df)}

    if domain_key == "知识产权":
        domain_config = DOMAINS[domain_key]
        core_fields = domain_config["core_fields"]

        group_field = next((c for c in ["系统划分区县", "所属区域", "区域", "地区"] if c in all_cols), None)
        patent_type_field = next((c for c in ["专利类型"] if c in all_cols), None)
        ipc_field = next((c for c in ["主分类号", "IPC分类号"] if c in all_cols), None)
        applicant_type_field = next((c for c in ["专利权人类型", "申请人类型"] if c in all_cols), None)
        applicant_name_field = next((c for c in ["专利权人名称", "申请人名称"] if c in all_cols), None)

        valid_patent_count = len(combined_df)
        totals["有效专利总量"] = valid_patent_count
        invention_patent_count = combined_df[combined_df[patent_type_field].astype(str).str.contains("发明", na=False)].shape[0] if patent_type_field else valid_patent_count
        totals["有效发明专利存量"] = invention_patent_count
        totals["发明专利占比"] = round(invention_patent_count / valid_patent_count * 100, 2) if valid_patent_count > 0 else 0
        totals["万人发明专利拥有量"] = round(invention_patent_count / 155, 2)

        totals["本月新增发明专利授权"] = len(dfs.get("f_auth", pd.DataFrame()))
        totals["本月PCT国际申请"] = len(dfs.get("f_pct", pd.DataFrame()))
        totals["本月失效专利"] = len(dfs.get("f_loss", pd.DataFrame()))
        totals["专利净增量"] = totals["本月新增发明专利授权"] - totals["本月失效专利"]

        top_data = {}
        wujiang_towns = ["吴江开发区", "汾湖", "盛泽", "太湖新城", "平望", "桃源", "七都", "震泽", "横扇", "同里"]
        if group_field:
            wujiang_df = combined_df[
                combined_df[group_field].astype(str).str.contains("|".join(wujiang_towns), na=False, case=False)
            ]
            town_counts = {}
            for town in wujiang_towns:
                count = wujiang_df[wujiang_df[group_field].astype(str).str.contains(f"^{town}$", na=False, case=False, regex=True)].shape[0]
                if count > 0:
                    town_counts[town] = count
            town_data = dict(sorted(town_counts.items(), key=lambda x: x[1], reverse=True))
            total_count = sum(town_counts.values())
            top_data = {"吴江区总计": total_count}
            top_data.update(town_data)
            totals["吴江区全区域专利总量"] = total_count
            totals["吴江区所有区域专利明细"] = f"吴江区总计{total_count}件（" + "、".join([f"{k}({v}件)" for k, v in town_data.items()]) + "）"
            totals["吴江区各镇街专利TOP3"] = "、".join([f"{k}({v}件)" for k, v in list(town_data.items())[:3]])

        top_data = {str(k): v for k, v in top_data.items()}
        type_data = combined_df[applicant_type_field].value_counts().to_dict() if applicant_type_field else {}
        type_data = {str(k): v for k, v in type_data.items()}
        class_data = combined_df[ipc_field].value_counts().head(10).to_dict() if ipc_field else {}
        class_data = {str(k): v for k, v in class_data.items()}
        name_data = combined_df[applicant_name_field].value_counts().head(10).to_dict() if applicant_name_field else {}
        name_data = {str(k): v for k, v in name_data.items()}
        reason_data = {}

        df_summary = pd.DataFrame({'专利数量': list(top_data.values())}, index=list(top_data.keys())) if top_data else pd.DataFrame()

        totals["top3_districts"] = "、".join(list(top_data.keys())[:3]) if top_data else "无"
        totals["top3_valid_districts"] = totals.get("吴江区各镇街专利TOP3", "无")
        totals["type_data"] = "、".join([f"{str(k)}({v}件)" for k, v in list(type_data.items())[:3]]) if type_data else "无"
        totals["top_applicant"] = "、".join([str(k) for k in list(name_data.keys())[:10]]) if name_data else "无"
        totals["top_ipc"] = "、".join([str(k) for k in list(class_data.keys())[:10]]) if class_data else "无"

    else:
        group_field = next((c for c in ["区域","地区","投诉地区","举报地区","系统划分区县"] if c in all_cols), None)
        type_field = next((c for c in ["类型","投诉类型","举报类型","问题类型","食品类别","企业类型"] if c in all_cols), None)
        reason_field = next((c for c in ["状态","处理状态","办理状态","问题类别"] if c in all_cols), None)
        class_field = next((c for c in ["主分类号","产品分类","行业分类"] if c in all_cols), None)
        name_field = next((c for c in ["企业名称","经营主体","被举报单位名称"] if c in all_cols), None)

        df_summary = pd.DataFrame({'数量': combined_df[group_field].value_counts().head(15)}) if group_field else pd.DataFrame()
        top_data = combined_df[group_field].value_counts().head(15).to_dict() if group_field else {}
        top_data = {str(k): v for k, v in top_data.items()}
        type_data = combined_df[type_field].value_counts().to_dict() if type_field else {}
        type_data = {str(k): v for k, v in type_data.items()}
        class_data = combined_df[class_field].value_counts().head(10).to_dict() if class_field else {}
        class_data = {str(k): v for k, v in class_data.items()}
        reason_data = combined_df[reason_field].value_counts().to_dict() if reason_field else {}
        reason_data = {str(k): v for k, v in reason_data.items()}
        name_data = combined_df[name_field].value_counts().head(10).to_dict() if name_field else {}
        name_data = {str(k): v for k, v in name_data.items()}

        if reason_field:
            totals["已办结"] = combined_df[reason_field].astype(str).str.contains("已办结").sum()
            totals["办结率"] = round(totals["已办结"]/len(combined_df)*100,1) if len(combined_df)>0 else 0

    totals["苏州大学关联数据"] = sum(1 for v in combined_df.astype(str).agg(' '.join, axis=1) if "苏州大学" in v)
    all_cols_info = {"列名": all_cols, "示例": combined_df.head(3).to_dict("records")}
    return totals, df_summary, top_data, type_data, class_data, reason_data, all_cols_info, name_data, dfs


def generate_domain_word(domain_key, ai_text, df_summary, top_data, type_data, class_data=None, name_data=None):
    domain = DOMAINS[domain_key]
    doc = Document()
    doc.styles['Normal'].font.name = u'仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

    domain_titles = {
        "知识产权": "吴江区知识产权统计简报",
        "投诉举报": "吴江区投诉举报分析简报",
        "食品安全": "吴江区食品安全监管分析简报",
        "医疗器械": "吴江区医疗器械监管检查情况分析简报"
    }
    domain_title = domain_titles.get(domain_key, "吴江区市场监管数据分析报告")

    title = doc.add_heading('', 0)
    run = title.add_run(domain_title)
    run.font.name = u'方正小标宋简体'
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('=' * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if domain_key == "知识产权":
        doc.add_heading('一、宏观发展概况', level=1)
        doc.add_paragraph(ai_text)
        doc.add_heading('二、核心指标动态', level=2)
        if not df_summary.empty:
            t_core = doc.add_table(rows=1, cols=2)
            t_core.style = 'Table Grid'
            t_core.rows[0].cells[0].text = '指标名称'
            t_core.rows[0].cells[1].text = '数值'
            for idx, row in df_summary.iterrows():
                row_cells = t_core.add_row().cells
                row_cells[0].text = str(idx)
                for i, col in enumerate(df_summary.columns):
                    row_cells[1].text = str(row[col])
        doc.add_heading('三、区域创新格局', level=2)
        if top_data:
            t_zone = doc.add_table(rows=1, cols=2)
            t_zone.style = 'Table Grid'
            t_zone.rows[0].cells[0].text = '区域'
            t_zone.rows[0].cells[1].text = '专利数量'
            for name, count in list(top_data.items())[:10]:
                row_cells = t_zone.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)
        doc.add_heading('四、创新主体结构', level=2)
        if type_data:
            t_type = doc.add_table(rows=1, cols=2)
            t_type.style = 'Table Grid'
            t_type.rows[0].cells[0].text = '主体类型'
            t_type.rows[0].cells[1].text = '数量'
            for name, count in list(type_data.items())[:10]:
                row_cells = t_type.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)
        if class_data:
            doc.add_heading('五、核心技术赛道', level=2)
            t_ipc = doc.add_table(rows=1, cols=2)
            t_ipc.style = 'Table Grid'
            t_ipc.rows[0].cells[0].text = 'IPC分类号'
            t_ipc.rows[0].cells[1].text = '专利数量'
            for name, count in list(class_data.items())[:10]:
                row_cells = t_ipc.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)
        if name_data:
            doc.add_heading('六、授权领军主体', level=2)
            t_name = doc.add_table(rows=1, cols=2)
            t_name.style = 'Table Grid'
            t_name.rows[0].cells[0].text = '主体名称'
            t_name.rows[0].cells[1].text = '授权数量'
            for name, count in list(name_data.items())[:10]:
                row_cells = t_name.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)
    else:
        doc.add_heading('一、数据概况', level=1)
        doc.add_paragraph(ai_text)

        if not df_summary.empty and top_data:
            doc.add_heading('二、区域分布明细', level=2)
            cols = len(df_summary.columns)
            t1 = doc.add_table(rows=1, cols=cols + 1)
            t1.style = 'Table Grid'
            headers = ['区域'] + list(df_summary.columns)
            for i, h in enumerate(headers):
                t1.rows[0].cells[i].text = h
            for idx, row in df_summary.iterrows():
                row_cells = t1.add_row().cells
                row_cells[0].text = str(idx)
                for i, col in enumerate(df_summary.columns):
                    row_cells[i + 1].text = str(row[col])

        if type_data:
            doc.add_heading('三、检查问题类型分布', level=2)
            t2 = doc.add_table(rows=1, cols=2)
            t2.style = 'Table Grid'
            t2.rows[0].cells[0].text = '类型'
            t2.rows[0].cells[1].text = '数量'
            for name, count in list(type_data.items())[:10]:
                row_cells = t2.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)

        if domain_key == "知识产权" and class_data:
            doc.add_heading('四、核心技术领域分布', level=2)
            t3 = doc.add_table(rows=1, cols=2)
            t3.style = 'Table Grid'
            t3.rows[0].cells[0].text = 'IPC分类号'
            t3.rows[0].cells[1].text = '专利数量'
            for name, count in list(class_data.items())[:10]:
                row_cells = t3.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)

        if name_data:
            doc.add_heading('四、核心主体TOP10' if domain_key != "知识产权" else '五、授权领军主体TOP10', level=2)
            t4 = doc.add_table(rows=1, cols=2)
            t4.style = 'Table Grid'
            t4.rows[0].cells[0].text = '主体名称'
            t4.rows[0].cells[1].text = '数量'
            for name, count in list(name_data.items())[:10]:
                row_cells = t4.add_row().cells
                row_cells[0].text, row_cells[1].text = str(name), str(count)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def generate_domain_pdf(domain_key, top_data, type_data, class_data, name_data):
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
    plt.rcParams['axes.unicode_minus'] = False
    figs = []

    if domain_key == "知识产权":
        if top_data:
            fig, ax = plt.subplots(figsize=(10,6))
            items = list(top_data.items())[:10]
            bars = ax.barh([str(k) for k,_ in items][::-1], [int(v) for _,v in items][::-1], color='#1565C0', edgecolor='white')
            for bar, val in zip(bars, [int(v) for _,v in items][::-1]):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=10)
            ax.set_title('吴江区各区域有效专利分布TOP10', fontsize=14, fontweight='bold')
            ax.tick_params(axis='x', rotation=0)
            ax.grid(axis='x', alpha=0.3)
            figs.append(fig)

        if type_data:
            fig, ax = plt.subplots(figsize=(10,6))
            keys = [str(k)[:12] for k in type_data.keys()]
            bars = ax.bar(keys, [int(v) for v in type_data.values()], color='#2E7D32', edgecolor='white')
            for bar, val in zip(bars, [int(v) for v in type_data.values()]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, str(val), ha='center', fontsize=9)
            ax.set_title('创新主体类型分布', fontsize=14, fontweight='bold')
            ax.tick_params(axis='x', rotation=45, labelsize=9)
            ax.grid(axis='y', alpha=0.3)
            figs.append(fig)

        if class_data:
            fig, ax = plt.subplots(figsize=(10,6))
            items = list(class_data.items())[:10]
            bars = ax.barh([str(k)[:15] for k,_ in items][::-1], [int(v) for _,v in items][::-1], color='#F57C00', edgecolor='white')
            for bar, val in zip(bars, [int(v) for _,v in items][::-1]):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=10)
            ax.set_title('核心技术领域IPC分布TOP10', fontsize=14, fontweight='bold')
            ax.grid(axis='x', alpha=0.3)
            figs.append(fig)

        if name_data:
            fig, ax = plt.subplots(figsize=(10,6))
            items = list(name_data.items())[:10]
            bars = ax.barh([str(k)[:12] for k,_ in items][::-1], [int(v) for _,v in items][::-1], color='#7B1FA2', edgecolor='white')
            for bar, val in zip(bars, [int(v) for _,v in items][::-1]):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=10)
            ax.set_title('授权领军主体TOP10', fontsize=14, fontweight='bold')
            ax.grid(axis='x', alpha=0.3)
            figs.append(fig)
    else:
        if top_data:
            fig, ax = plt.subplots(figsize=(10,6))
            items = list(top_data.items())[:10]
            bars = ax.barh([str(k) for k,_ in items][::-1], [int(v) for _,v in items][::-1], color='#1E88E5', edgecolor='white')
            for bar, val in zip(bars, [int(v) for _,v in items][::-1]):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=10)
            ax.set_title('吴江区各镇/街道检查数量TOP10', fontsize=14, fontweight='bold')
            ax.tick_params(axis='x', rotation=0)
            ax.grid(axis='x', alpha=0.3)
            figs.append(fig)

        if type_data:
            fig, ax = plt.subplots(figsize=(10,6))
            keys = [str(k)[:12] for k in type_data.keys()]
            bars = ax.bar(keys, [int(v) for v in type_data.values()], color='#E53935', edgecolor='white')
            for bar, val in zip(bars, [int(v) for v in type_data.values()]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, str(val), ha='center', fontsize=9)
            ax.set_title('检查问题类型分布', fontsize=14, fontweight='bold')
            ax.tick_params(axis='x', rotation=45, labelsize=9)
            ax.grid(axis='y', alpha=0.3)
            figs.append(fig)

        if type_data:
            fig, ax = plt.subplots(figsize=(8,6))
            labels = [str(k)[:8] for k in list(type_data.keys())[:6]]
            values = [int(v) for v in list(type_data.values())[:6]]
            colors = ['#1E88E5', '#43A047', '#FB8C00', '#E53935', '#7B1FA2', '#FDD835']
            wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors[:len(values)])
            for autotext in autotexts:
                autotext.set_fontsize(10)
            ax.set_title('监管对象类型分布', fontsize=14, fontweight='bold')
            figs.append(fig)

        if name_data:
            fig, ax = plt.subplots(figsize=(10,6))
            items = list(name_data.items())[:10]
            bars = ax.barh([str(k)[:12] for k,_ in items][::-1], [int(v) for _,v in items][::-1], color='#FB8C00', edgecolor='white')
            for bar, val in zip(bars, [int(v) for _,v in items][::-1]):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(val), va='center', fontsize=10)
            ax.set_title('涉事主体TOP10', fontsize=14, fontweight='bold')
            ax.grid(axis='x', alpha=0.3)
            figs.append(fig)

    if not figs:
        fig, ax = plt.subplots()
        ax.text(0.5,0.5,"暂无数据",ha='center')
        figs.append(fig)

    pdf_io = io.BytesIO()
    with PdfPages(pdf_io) as pdf:
        for fig in figs:
            pdf.savefig(fig, bbox_inches='tight')
            plt.close(fig)
    pdf_io.seek(0)
    return pdf_io


def generate_domain_excel(domain_key, totals, df_summary, top_data, type_data, class_data, name_data=None, dfs=None):
    excel_io = io.BytesIO()
    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
        has_data = False

        overview_data = [(str(k), str(v)) for k, v in totals.items()]
        overview_df = pd.DataFrame(overview_data, columns=['指标名称', '数值'])
        overview_df.to_excel(writer, sheet_name='数据总览', index=False)
        has_data = True

        if not df_summary.empty:
            df_summary.reset_index().rename(columns={'index': '区域名称'}).to_excel(writer, sheet_name='区域分布明细', index=False)
            has_data = True

        if type_data:
            type_df = pd.DataFrame([(str(k), str(v)) for k, v in type_data.items()], columns=['类型', '数量'])
            type_df.to_excel(writer, sheet_name='类型分布统计', index=False)
            has_data = True
        
        if domain_key == "知识产权" and class_data:
            class_df = pd.DataFrame([(str(k), str(v)) for k, v in class_data.items()], columns=['IPC分类号', '数量'])
            class_df.to_excel(writer, sheet_name='技术领域分布', index=False)
            has_data = True

        if name_data:
            name_df = pd.DataFrame([(str(k), str(v)) for k, v in name_data.items()], columns=['主体名称', '数量'])
            name_df.to_excel(writer, sheet_name='核心主体TOP10', index=False)
            has_data = True

        if "已办结" in totals and "办结率" in totals:
            status_df = pd.DataFrame([
                {"状态": "已办结", "数量": str(totals.get("已办结", 0))},
                {"状态": "办结率", "数量": f"{str(totals.get('办结率', 0))}%"}
            ])
            status_df.to_excel(writer, sheet_name='处理状态统计', index=False)
            has_data = True

        if dfs:
            for key, df in dfs.items():
                df.to_excel(writer, sheet_name='原始数据明细', index=False)
                has_data = True
        
        if not has_data:
            pd.DataFrame({'提示': ['无数据']}).to_excel(writer, sheet_name='汇总')
    excel_io.seek(0)
    return excel_io


col_left, col_right = st.columns([1, 2])

with col_left:
    st.markdown('<div class="flat-title">📁 数据上传与领域选择</div>', unsafe_allow_html=True)

    st.markdown("##### 🏷️ 选择分析领域")
    preset_domains = [
        ("知识产权", "📈", "专利申请、授权、失效、PCT等数据分析"),
        ("投诉举报", "📞", "12345投诉举报数据分析"),
        ("食品安全", "🍱", "食品安全监管、抽检、处罚数据分析"),
        ("医疗器械", "🏥", "医疗器械生产、经营、监管数据分析"),
        ("自定义", "📋", "自定义上传数据并分析")
    ]

    domain_options = [f"{d[1]} {d[0]}" for d in preset_domains]
    domain_keys = [d[0] for d in preset_domains]

    selected_domain_display = st.selectbox(
        "选择市场监管分析领域",
        domain_options,
        index=domain_keys.index(st.session_state.selected_domain) if st.session_state.selected_domain in domain_keys else 0,
        label_visibility="collapsed"
    )

    selected_idx = domain_options.index(selected_domain_display)
    st.session_state.selected_domain = domain_keys[selected_idx]

    domain_desc = ""
    for d in preset_domains:
        if d[0] == st.session_state.selected_domain:
            domain_desc = d[2]
            break
    st.markdown(f"*{domain_desc}*")

    st.markdown("##### 📤 上传业务清单")
    uploaded_files = {}
    
    if st.session_state.selected_domain == "知识产权":
        ip_files = [
            {"key": "f_valid", "label": "《有效专利清单》"},
            {"key": "f_pct", "label": "《PCT申请清单》"},
            {"key": "f_auth", "label": "《发明授权清单》"},
            {"key": "f_loss", "label": "《失效专利清单》"}
        ]
        for f in ip_files:
            uploaded_files[f["key"]] = st.file_uploader(f["label"], type=["xlsx", "xls", "csv"], key=f"file_{f['key']}")
    elif st.session_state.selected_domain == "自定义":
        uploaded = st.file_uploader("上传数据文件（支持多文件）", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="custom_multi_files")
        if uploaded:
            for i, f in enumerate(uploaded):
                uploaded_files[f"custom_{i}"] = f
    else:
        # 投诉举报、食品安全、医疗器械等 - 直接多文件上传
        uploaded = st.file_uploader("上传业务数据文件（支持多文件）", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="biz_files")
        if uploaded:
            for i, f in enumerate(uploaded):
                uploaded_files[f"biz_{i}"] = f

    # 显示已上传文件状态
    file_count = sum(1 for v in uploaded_files.values() if v is not None)
    if file_count > 0:
        st.success(f"✅ 已上传 {file_count} 个文件")
    else:
        st.warning("⏳ 请上传数据文件...")

    st.markdown("##### ⚙️ 大模型引擎")
    selected_model_input = st.selectbox(
        "选择大模型",
        ("豆包大模型 (Volcengine)", "DeepSeek"),
        label_visibility="collapsed"
    )

    st.markdown("<br>", unsafe_allow_html=True)
    
    # 根据文件是否上传决定按钮是否可用
    start_btn = False
    if file_count == 0:
        st.button("🚀 开始自动化分析", use_container_width=True, type="primary", disabled=True)
        st.info("💡 请先上传数据文件后，再点击开始")
    else:
        start_btn = st.button("🚀 开始自动化分析", use_container_width=True, type="primary")

if start_btn:
    uploaded_count = sum(1 for v in uploaded_files.values() if v is not None)
    if uploaded_count == 0:
        st.warning("⚠️ 提示：请先上传数据文件，等待文件显示在页面后再点击开始！")
        st.stop()
    
    # 验证文件是否有效
    valid_files = {}
    for k, v in uploaded_files.items():
        if v is not None and hasattr(v, 'name'):
            valid_files[k] = v
    
    if len(valid_files) == 0:
        st.warning("⚠️ 文件上传不完整，请重新选择文件后等待显示后再点击开始！")
        st.stop()
    
    st.session_state.selected_model = selected_model_input
    st.session_state.started = True
    st.session_state.processed = False
    st.session_state.uploaded_files = valid_files
    st.rerun()

with col_right:
    st.markdown('<div class="flat-title">🎛️ 智能分析控制台</div>', unsafe_allow_html=True)

    if not st.session_state.started:
        st.info("👈 请在左侧选择领域、上传数据，选择分析引擎，点击开始进行智能分析。")

    if st.session_state.started:
        if not st.session_state.processed:
            domain = DOMAINS[st.session_state.selected_domain]
            with st.status("智能体全链路运作中...", expanded=True) as status:
                st.write("📥 [1/4] 正在解析表格并执行多维度分析...")

                try:
                    totals, df_summary, top_data, type_data, class_data, reason_data, all_cols_info, name_data, dfs = process_domain_data(
                        st.session_state.selected_domain, st.session_state.uploaded_files)
                    all_cols = all_cols_info.get("列名", []) if all_cols_info else []
                except Exception as e:
                    status.update(label="处理异常中断！", state="error", expanded=True)
                    st.error(f"🚨 【格式不规范提醒】您上传的文件内容不符合系统要求！错误详情: {str(e)}")
                    st.stop()

                st.write(f"✍️ [2/4] 正在唤醒 {st.session_state.selected_model} 撰写分析简报...")

                if st.session_state.selected_model == "豆包大模型 (Volcengine)":
                    c_key = VOLC_API_KEY
                    c_url = VOLC_BASE_URL
                    c_name = VOLC_MODEL_NAME
                else:
                    c_key = DEEPSEEK_API_KEY
                    c_url = DEEPSEEK_BASE_URL
                    c_name = DEEPSEEK_MODEL_NAME

                domain_key = st.session_state.selected_domain
                
                if domain_key == "知识产权":
                    domain_config = DOMAINS[domain_key]
                    try:
                        format_params = {
                            "吴江区全区域专利总量": str(totals.get("吴江区全区域专利总量", 0)),
                            "有效发明专利存量": str(totals.get("有效发明专利存量", 0)),
                            "本月新增发明专利授权": str(totals.get("本月新增发明专利授权", 0)),
                            "本月PCT国际申请": str(totals.get("本月PCT国际申请", 0)),
                            "本月失效专利": str(totals.get("本月失效专利", 0)),
                            "吴江区所有区域专利明细": totals.get("吴江区所有区域专利明细", "无"),
                            "top_applicant": totals.get("top_applicant", "无"),
                            "top_ipc": totals.get("top_ipc", "无")
                        }
                        prompt = domain_config["prompt_template"].format(**format_params)
                    except Exception as e:
                        print(f"Prompt拼接出错: {str(e)}")
                        prompt = f"你是吴江区市场监管局知识产权局高级政务分析师，基于本次上传的{str(totals.get('文件数', 0))}个文件，共{str(totals.get('总记录数', 0))}条专利数据，撰写300-400字的专业分析简报，严格遵循苏州市官方简报标准，所有内容基于真实数据，无编造、无套话。"
                else:
                    prompt_parts = ["你是政务数据分析师，只基于上传真实数据写300-400字简报，不编造、不套话。"]
                    prompt_parts.append(f"概况：{str(totals.get('文件数', 0))}个文件，{str(totals.get('总记录数', 0))}条，字段：{', '.join([str(c) for c in all_cols])}")
                    if class_data: prompt_parts.append(f"违规类型：{', '.join([f'{str(k)}({str(v)})' for k,v in class_data.items()])}")
                    if type_data: prompt_parts.append(f"举报类型：{', '.join([f'{str(k)}({str(v)})' for k,v in type_data.items()])}")
                    if top_data: prompt_parts.append(f"区域：{', '.join([f'{str(k)}({str(v)})' for k,v in list(top_data.items())[:3]])}")
                    if name_data: prompt_parts.append(f"主体TOP3：{', '.join([str(k) for k in list(name_data.keys())[:3]])}")
                    prompt = "。".join(prompt_parts) + "。要求：数据概况、特征、风险、建议，政府口吻，只写数据内内容。"

                try:
                    client = OpenAI(api_key=c_key, base_url=c_url)
                    response = client.chat.completions.create(
                        model=c_name,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3
                    )
                    final_text = response.choices[0].message.content
                except openai.APIConnectionError:
                    status.update(label="网络连接失败", state="error", expanded=True)
                    st.error(f"🚨 **【网络连接失败】**无法连接到 {st.session_state.selected_model} 的服务器！请检查网络或代理设置。")
                    st.stop()
                except openai.AuthenticationError:
                    status.update(label="API 鉴权失败", state="error", expanded=True)
                    st.error(f"🚨 **【API Key 无效或过期】**：您提供的 {st.session_state.selected_model} 密钥无法通过验证。")
                    st.stop()
                except Exception as e:
                    status.update(label="API 调用异常", state="error", expanded=True)
                    st.error(f"🚨 **【调用失败】**：{str(e)}")
                    st.stop()

                st.write("📑 [3/4] 正在生成 Word、Excel、PDF 多形态报告...")
                st.session_state.word_file = generate_domain_word(st.session_state.selected_domain, final_text, df_summary, top_data, type_data, class_data, name_data)
                st.session_state.excel_file = generate_domain_excel(st.session_state.selected_domain, totals, df_summary, top_data, type_data, class_data, name_data, dfs)
                st.session_state.pdf_file = generate_domain_pdf(
                    st.session_state.selected_domain, top_data, type_data, class_data, name_data)
                st.session_state.final_text = final_text
                st.session_state.domain_title = domain["title"]

                st.session_state.processed = True
                st.write("✅ [4/4] 全部打包完毕！")
                status.update(label="处理完成", state="complete", expanded=False)
                st.rerun()

        if st.session_state.processed:
            st.success("🎉 数据深加工完成！")

            btn_col1, btn_col2, btn_col3 = st.columns(3)
            with btn_col1:
                st.download_button("📝 下载分析简报", data=st.session_state.word_file,
                                   file_name=f"{st.session_state.selected_domain}_分析简报.docx", use_container_width=True)
            with btn_col2:
                st.download_button("📊 下载数据底稿", data=st.session_state.excel_file,
                                   file_name=f"{st.session_state.selected_domain}_数据底稿.xlsx", use_container_width=True)
            with btn_col3:
                st.download_button("📈 下载可视化大屏", data=st.session_state.pdf_file,
                                   file_name=f"{st.session_state.selected_domain}_分析大屏.pdf", use_container_width=True)

            st.markdown('<div class="flat-title" style="margin-top: 30px;">📄 AI 分析报告正文</div>',
                        unsafe_allow_html=True)
            st.info(st.session_state.final_text)
